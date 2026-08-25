"""Build a scenario-level validation summary table from report_summary.tsv."""

import argparse
from pathlib import Path

import polars as pl

# ---------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------

# Column names below must match the schema produced by
# simace.analysis.gather.extract_metrics (report_summary.tsv); the
# observed-metric columns are defined in
# simace.analysis.report_schema.REPORT_SUMMARY_REGISTRY.

# Columns representing true simulation parameters
# (constant within each scenario)
TRUE_COLS = ["N", "A1", "C1", "E1", "A2", "C2", "E2", "rA", "rC"]

# Observed numeric columns to summarize as mean ± standard deviation
NUMERIC_COLS = [
    "variance_A1",
    "variance_C1",
    "variance_E1",
    "variance_A2",
    "variance_C2",
    "variance_E2",
    "observed_rA",
    "observed_rC",
    "observed_rE",
    "falconer_h2_trait1",
    "falconer_h2_trait2",
    "mz_twin_liability1_corr",
    "mz_twin_liability2_corr",
    "dz_sibling_liability1_corr",
    "dz_sibling_liability2_corr",
    "half_sib_liability1_corr",
    "mate_corr_liability1",
    "mate_corr_liability2",
    "parent_offspring_liability1_slope",
    "parent_offspring_liability2_slope",
]

# Column order for the final paper-ready table
PAPER_COL_ORDER = [
    "scenario",
    "N",
    "A1",
    "falconer_h2_trait1",
    "variance_A1",
    "A2",
    "falconer_h2_trait2",
    "variance_A2",
    "C1",
    "variance_C1",
    "C2",
    "variance_C2",
    "rA",
    "observed_rA",
    "mz_twin_liability1_corr",
    "dz_sibling_liability1_corr",
    "mz_twin_liability2_corr",
    "dz_sibling_liability2_corr",
    "mate_corr_liability1",
    "mate_corr_liability2",
]


# ---------------------------------------------------------------------
# Table construction
# ---------------------------------------------------------------------


def build_paper_table(df: pl.DataFrame) -> pl.DataFrame:
    """Collapse replication-level results into one row per scenario.

    Each numeric column is summarized as ``mean ± standard deviation``.
    """
    rows = []

    for (scenario,), group in df.group_by("scenario", maintain_order=True):
        row = {"scenario": scenario}

        for col in NUMERIC_COLS:
            if col not in group.columns:
                continue
            mean = group[col].mean()
            if mean is None:
                # Every rep is null for this scenario (e.g. observed_rC where
                # rC == 0) — leave the cell blank rather than formatting a None.
                row[col] = None
                continue
            sd = group[col].std()
            if sd is not None:
                row[col] = f"{mean:.4f} ± {sd:.4f}"
            else:
                row[col] = f"{mean:.4f}"

        rows.append(row)

    paper_df = pl.DataFrame(rows)

    # Attach true simulation parameters (identical within each scenario)
    true_present = [c for c in TRUE_COLS if c in df.columns]
    if true_present:
        true_vals = df.group_by("scenario", maintain_order=True).agg(pl.col(c).first() for c in true_present)
        paper_df = paper_df.join(true_vals, on="scenario", how="left", maintain_order="left")

    # Enforce predefined column order
    ordered_cols = [c for c in PAPER_COL_ORDER if c in paper_df.columns]
    return paper_df.select(ordered_cols)


# ---------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------


def save_word_table(paper_df: pl.DataFrame, out_path: Path) -> Path:
    """Save the summary table as a Word document using python-docx."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENTATION
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise ImportError("--word requires python-docx; install with `pip install python-docx`.") from exc

    doc = Document()

    # Page setup: A4 landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Table title
    title = doc.add_heading("Table 1 — Validation summary (mean ± SD across replications)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph("")

    # Create table
    n_cols = paper_df.width
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(paper_df.columns):
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)

    # Data rows
    for row in paper_df.iter_rows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            text = "" if val is None else str(val)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(8)

    word_file = out_path / "table_1.docx"
    doc.save(word_file)
    return word_file


# ---------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------


def main(summary_tsv: str, out_dir: str = "paper", word: bool = False) -> None:
    """Read a validation summary TSV and write a scenario-level CSV (and optional DOCX)."""
    summary_path = Path(summary_tsv)
    if not summary_path.exists():
        raise FileNotFoundError(f"Summary TSV not found: {summary_path.resolve()}")

    # infer_schema_length=None: the per-generation dict strings in A1/C1/E1 can
    # land past polars' default 100-row inference window.
    df = pl.read_csv(summary_path, separator="\t", null_values=["NA", ""], infer_schema_length=None)

    if "scenario" not in df.columns:
        raise ValueError("Expected a 'scenario' column in the input file.")

    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    # Build scenario-level summary table
    paper_df = build_paper_table(df)
    paper_file = out_path / "table_1.csv"
    paper_df.write_csv(paper_file)

    print("\n=== Table 1 (mean ± SD across replications) ===")
    with pl.Config(tbl_rows=-1, tbl_cols=-1):
        print(paper_df)

    print("\nSaved files:")
    print(f"  {paper_file}")

    if word:
        # python-docx is an optional dependency; only required for --word
        word_file = save_word_table(paper_df, out_path)
        print(f"  {word_file}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build a scenario-level validation summary table.")
    parser.add_argument(
        "--summary",
        required=True,
        help="Path to report_summary.tsv.",
    )
    parser.add_argument(
        "--out_dir",
        default="paper",
        help="Output directory (default: paper/).",
    )
    parser.add_argument(
        "--word",
        action="store_true",
        help="Also export table_1.docx (requires python-docx).",
    )

    args = parser.parse_args()
    main(args.summary, args.out_dir, word=args.word)
